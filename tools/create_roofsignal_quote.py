#!/usr/bin/env python3
"""Generate a RoofSignal quotation from the fixed Koorn visual template.

Usage:
  python tools/create_roofsignal_quote.py quote.json --output output/offerte.docx

Customer JSON files are intentionally not committed. The visual template and
this generator are the single source of truth for RoofSignal quotations.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = ROOT / "templates" / "roofsignal-offerte-template.docx"


def money(value: float) -> str:
    formatted = f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"€ {formatted}"


def replace_paragraph(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_cell(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        replace_paragraph(extra, "")


def required(data: dict, name: str):
    value = data.get(name)
    if value is None or value == "":
        raise ValueError(f"Verplicht offerteveld ontbreekt: {name}")
    return value


def build(data: dict, template: Path, output: Path) -> None:
    doc = Document(template)

    # Fixed paragraph slots from the canonical two-page quotation template.
    replace_paragraph(doc.paragraphs[2], required(data, "package_title"))
    replace_paragraph(doc.paragraphs[3], data.get("status_line", "CONCEPTOFFERTE • NOG NIET VERZONDEN"))
    replace_paragraph(doc.paragraphs[5], required(data, "customer_question"))

    scope = required(data, "scope_bullets")
    if len(scope) != 8:
        raise ValueError("scope_bullets moet exact 8 regels bevatten voor vaste paginalay-out")
    for paragraph, text in zip(doc.paragraphs[7:15], scope):
        replace_paragraph(paragraph, text)

    replace_paragraph(doc.paragraphs[16], data.get("price_note", "Reisafstand en bereikbaarheid zijn verwerkt volgens de RoofSignal-prijsafspraken."))

    working = required(data, "working_bullets")
    if len(working) != 6:
        raise ValueError("working_bullets moet exact 6 regels bevatten voor vaste paginalay-out")
    for paragraph, text in zip(doc.paragraphs[19:25], working):
        replace_paragraph(paragraph, text)

    replace_paragraph(doc.paragraphs[26], required(data, "delivery_text"))
    replace_paragraph(doc.paragraphs[28], required(data, "payment_validity_text"))
    replace_paragraph(doc.paragraphs[30], data.get("acceptance_text", "Door ondertekening geeft opdrachtgever akkoord op de beschreven scope, investering en uitgangspunten."))

    # These source paragraphs were centred in an early Koorn working file.
    # Force regular left-aligned body copy so every section starts cleanly
    # below its heading in Word as well as in the generated PDF.
    for index in (26, 28, 30):
        paragraph = doc.paragraphs[index]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)

    # LibreOffice can collapse the source template's keep-with-next heading
    # geometry on page two. Fix the rhythm explicitly while retaining the
    # template's font, colour and size.
    for index in (18, 25, 27, 29):
        paragraph = doc.paragraphs[index]
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)

    metadata = doc.tables[0]
    meta_values = [
        ("Offerte nr.", required(data, "quote_number"), "Datum", required(data, "quote_date")),
        ("Geldig tot", required(data, "valid_until"), "Pakket", required(data, "package")),
        ("Klant", required(data, "customer_name"), "E-mail", data.get("customer_email") or "Niet opgegeven"),
        ("Klantadres", required(data, "customer_address"), "Inspectie", required(data, "inspection_address")),
    ]
    for row, values in zip(metadata.rows, meta_values):
        for cell, value in zip(row.cells, values):
            replace_cell(cell, str(value))
    for row in metadata.rows:
        for index in (0, 2):
            for run in row.cells[index].paragraphs[0].runs:
                run.font.size = Pt(9)

    service_amount = float(required(data, "service_amount_ex_vat"))
    travel_amount = float(data.get("travel_amount_ex_vat", 0))
    subtotal = service_amount + travel_amount
    vat_rate = float(data.get("vat_rate", 21))
    vat = round(subtotal * vat_rate / 100, 2)
    total = subtotal + vat

    price = doc.tables[1]
    price_rows = [
        (required(data, "service_label"), "1", money(service_amount)),
        (required(data, "travel_label"), "", money(travel_amount)),
        ("Subtotaal", "", money(subtotal)),
        (f"{vat_rate:g}% btw", "", money(vat)),
    ]
    for row, values in zip(price.rows[1:], price_rows):
        for cell, value in zip(row.cells, values):
            replace_cell(cell, str(value))

    total_table = doc.tables[2]
    replace_cell(total_table.cell(0, 0), "Totaal inclusief btw")
    replace_cell(total_table.cell(0, 1), money(total))

    acceptance = doc.tables[3]
    replace_cell(acceptance.cell(1, 0), f"Naam: {required(data, 'signer_customer')}")
    replace_cell(acceptance.cell(1, 1), f"Naam: {data.get('signer_roofsignal', 'F.J. Joosten')}")

    doc.core_properties.title = f"Offerte {data['customer_name']} - {data['package']}"
    doc.core_properties.subject = data["package_title"]
    doc.core_properties.author = "F.J. Joosten - RoofSignal"
    doc.core_properties.comments = "Gegenereerd met het vaste RoofSignal-offertemodel"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    build(data, args.template, args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
